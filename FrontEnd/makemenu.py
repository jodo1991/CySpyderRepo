import os
from tkinter import filedialog, Menu, messagebox
import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

class MainMenu:
    def __init__(self, tkapp):
        self.tkapp = tkapp
        self.saveFilename = None

    def createMenu(self):
        topmenu = Menu(self.tkapp)

        filemenu = Menu(topmenu, tearoff=0)
        filemenu.add_command(label="Open", command=self.openMenu)
        filemenu.add_command(label="Save", command=lambda: self.saveMenu("Save"))
        filemenu.add_command(label="Save As", command=lambda: self.saveMenu("Save As"))
        filemenu.add_separator()

        filemenu.add_command(label="Quit", command=self.tkapp.quit)
        topmenu.add_cascade(label="File", menu=filemenu)

        helpmenu = Menu(topmenu, tearoff=0)
        helpmenu.add_command(label="About", command=self.aboutMenu)
        topmenu.add_cascade(label="Help", menu=helpmenu)

        # show menu
        self.tkapp.config(menu=topmenu)


    def openMenu(self):
        # set initial directory to savedScans folder
        try:
            self.tkapp.frames['SearchFrame'].deletesearch()
        except AttributeError:
            pass

        path = os.path.join(os.getcwd(), "Sessions")
        if not os.path.exists(path):
            os.makedirs(path)
        path = filedialog.askopenfilename(initialdir=path)
        try:
            with open(path, "r") as fin:
                self.tkapp.set_file(json.load(fin))

            self.tkapp.frames['SearchFrame'].helper.hidesearch()
            self.tkapp.frames['StartFrame'].start('open file')

        except FileNotFoundError:
            pass

    def saveMenu(self, saveType):
        # create main directory and subdir(current date) if not made already
        path = os.getcwd() + "/Sessions/" + str(datetime.date.today())
        if not os.path.exists(path):
            os.makedirs(path)

        if saveType == "Save" and self.saveFilename is not None:
            if messagebox.askyesnocancel("Overwrite File?", "Overwrite {}?".format(self.saveFilename)):
                with open(os.path.join(path, self.saveFilename), 'w') as file:
                    json.dump(self.tkapp.get_file(), file)
        else:
            # get a filename from the user or default to current time
            currentTime = datetime.datetime.now().strftime("%H_%M_%S")
            filename = filedialog.asksaveasfilename(defaultextension="txt", initialdir=path, initialfile=currentTime)
            if filename:
                self.saveFilename = filename
                with open(filename, 'w') as f:
                    json.dump(self.tkapp.get_file(), f)

    def aboutMenu(self):
        messagebox.showinfo("About", "Developed by: Chase Brown and Joseph Dodson\n""Refer to README for additional info.")


class ArticleMenu:

    def __init__(self, window, textbox, n):

        self.article = {"uri":n[2], "title":n[1], "author":n[3], "date":n[0], "body":n[4]}
        self.tb = textbox
        topmenu = Menu(window, tearoff=0)
        window.config(menu=topmenu)

        filemenu = Menu(topmenu, tearoff=0)
        topmenu.add_cascade(label="File", menu=filemenu)

        exportmenu = Menu(filemenu, tearoff=0)
        exportmenu.add_command(label="pdf", command=lambda: self.exportarticle("pdf"))
        exportmenu.add_command(label="docx", command=lambda: self.exportarticle("docx"))
        filemenu.add_cascade(label="Export", menu=exportmenu)

        settingsmenu = Menu(topmenu, tearoff=0)
        topmenu.add_cascade(label="Settings", menu=settingsmenu)

        fontsizemenu = Menu(settingsmenu, tearoff=0)
        fontsizemenu.add_command(label="Default", command=lambda: self.changefontsize(" 14"))
        fontsizemenu.add_command(label="16", command=lambda: self.changefontsize(" 16"))
        fontsizemenu.add_command(label="20", command=lambda: self.changefontsize(" 20"))
        fontsizemenu.add_command(label="24", command=lambda: self.changefontsize(" 24"))
        fontsizemenu.add_command(label="30", command=lambda: self.changefontsize(" 30"))
        settingsmenu.add_cascade(label="Change Font Size", menu=fontsizemenu)

        fontnamemenu = Menu(settingsmenu, tearoff=0)
        fontnamemenu.add_command(label="Times New Roman", command=lambda: self.changefontname("Times "))
        fontnamemenu.add_command(label="Courier", command=lambda: self.changefontname("Courier "))
        fontnamemenu.add_command(label="Helvetica", command=lambda: self.changefontname("Helvetica "))
        settingsmenu.add_cascade(label="Change Font Name", menu=fontnamemenu)

    def exportarticle(self, filetype):

        path = os.path.normpath(os.getcwd() + os.sep + os.pardir + '\Exports\\')
        if not os.path.exists(path):
            os.makedirs(path)

        currentTime = datetime.datetime.now().strftime("%H_%M_%S")

        defaultfilename = ' '.join(self.article['title'].split()[0:4]) + currentTime
        filename = filedialog.asksaveasfilename(initialdir=path, initialfile=defaultfilename)
        path = os.path.join(path, filename)

        if len(filename) > 1:
            if filetype == 'pdf':
                style = getSampleStyleSheet()
                sty = getSampleStyleSheet()['Title']
                sty.leading = 80
                pdf = SimpleDocTemplate(path + '.pdf', pagesize=letter)

                fontsize = "40" if (len(self.article['title'])) < 70 else "34"
                article = []
                article.append(Spacer(0, inch))
                article.append(Paragraph("<para fontSize={} spaceAfter=40>{}</para>".format(fontsize, self.article['title']), sty))
                article.append(
                    Paragraph("<para fontSize=20 align=center>Author: {}</para>".format(self.article['author']), style['Normal']))
                article.append(Spacer(0, inch*.35))
                article.append(
                    Paragraph("<para fontSize=15 align=center spaceAfter=70>Date Published: {}</para>".format(self.article['date']),
                              style["Normal"]))
                article.append(
                    Paragraph("<para alignment=center><link fontSize=12 textColor=blue "
                              "href={}><u>{}</u></link></para>".format(self.article['uri'], self.article['uri']), style['Normal']))
                article.append(PageBreak())

                for para in self.article['body'].split('\n\n'):
                    article.append(Paragraph(para, style['Normal']))
                    article.append(Spacer(.25 * inch, .25 * inch))
                pdf.build(article)

            if filetype == 'docx':
                document = Document()

                document.add_paragraph('\n\n\n\n')

                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.line_spacing = 1.5
                p = p.add_run(self.article['title'] + '\n')
                p.font.size = Pt(40) if len(self.article['title']) < 70 else Pt(32)
                p.font.bold = True

                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run("Author: {}".format(self.article['author'])). font.size = Pt(20)

                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run("Date Published: {}\n\n".format(self.article['date'])).font.size = Pt(15)

                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p = p.add_run(self.article['uri'])
                p.font.color.rgb = RGBColor(0x00,0x00,0xFF)
                p.font.underline = True
                document.add_page_break()

                for para in self.article['body'].split('\n\n'):
                    document.add_paragraph(para.replace('\n', ' '))
                document.save(path + '.docx')

    def changefontsize(self, fontsize):
        self.tb.config(font=self.tb.cget('font').split()[0] + fontsize)

    def changefontname(self, fontname):
        self.tb.config(font=fontname + self.tb.cget('font').split()[1])
